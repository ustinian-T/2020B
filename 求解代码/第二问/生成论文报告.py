from __future__ import annotations
import csv,json
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm,Pt,RGBColor
ROOT=Path(__file__).resolve().parents[2]; RES=ROOT/'求解代码'/'第二问'/'结果输出'; OUT=ROOT/'求解报告'/'第二问模型建立与求解及结果分析.docx'
SUMMARY=json.loads((RES/'求解摘要.json').read_text(encoding='utf-8'))
def rows(n):
 with (RES/n).open(encoding='utf-8-sig',newline='') as f:return list(csv.DictReader(f))
MET=rows('第四关蒙特卡洛指标对比.csv'); GAM=rows('第四关Gamma灵敏性分析.csv'); STO=rows('第四关沙暴概率灵敏性分析.csv')
BLUE=RGBColor(31,78,121)
def ft(r,n=10.5,b=False,e='宋体',a='Times New Roman',c=None):
 r.font.name=a;r._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),e);r.font.size=Pt(n);r.bold=b
 if c:r.font.color.rgb=c
def hd(d,s,l=1):
 p=d.add_paragraph(style=f'Heading {l}');ft(p.add_run(s),{1:15,2:13,3:11.5}[l],True,'黑体','Arial',BLUE)
def tx(d,s,indent=True):
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;p.paragraph_format.line_spacing=1.5;p.paragraph_format.space_after=Pt(5)
 if indent:p.paragraph_format.first_line_indent=Pt(21)
 ft(p.add_run(s))
def eq(d,s,n):
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;ft(p.add_run(f'{s}                         ({n})'),10.5,False,'宋体','Cambria Math')
def tb(d,h,rs,size=8.5):
 t=d.add_table(rows=1,cols=len(h));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for i,x in enumerate(h):t.rows[0].cells[i].text=str(x);sh=OxmlElement('w:shd');sh.set(qn('w:fill'),'EAF0F6');t.rows[0].cells[i]._tc.get_or_add_tcPr().append(sh)
 for r in rs:
  c=t.add_row().cells
  for i,x in enumerate(r):c[i].text=str(x)
 for ri,r in enumerate(t.rows):
  for c in r.cells:
   for p in c.paragraphs:
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(0)
    for z in p.runs:ft(z,size,ri==0,'黑体' if ri==0 else '宋体')
 d.add_paragraph()
def bl(d,title,s):
 p=d.add_paragraph(style='List Bullet');ft(p.add_run(title),10.5,True,'黑体');ft(p.add_run(s))
def setup(d):
 s=d.sections[0];s.page_width=Cm(21);s.page_height=Cm(29.7);s.top_margin=s.bottom_margin=Cm(2.4);s.left_margin=s.right_margin=Cm(2.5)
 rpr=d.styles['Normal']._element.get_or_add_rPr();rpr.get_or_add_rFonts().set(qn('w:eastAsia'),'宋体')
 p=s.header.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;ft(p.add_run('2020B 穿越沙漠——第二问模型建立与求解及结果分析'),9,c=RGBColor(100,100,100))
 p=s.footer.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('第 ');ft(r,9);b=OxmlElement('w:fldChar');b.set(qn('w:fldCharType'),'begin');i=OxmlElement('w:instrText');i.text=' PAGE ';e=OxmlElement('w:fldChar');e.set(qn('w:fldCharType'),'end');r._r.extend((b,i,e));ft(p.add_run(' 页'),9)
def build():
 d=Document();setup(d);d.core_properties.title='第二问模型建立与求解及结果分析'
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(80);ft(p.add_run('第二问：模型建立与求解及结果分析'),22,True,'黑体','Arial')
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;ft(p.add_run('自适应鲁棒动态规划 · 非前视情景树 · 预算不确定集 · Monte Carlo检验'),12,False,'黑体',c=BLUE)
 tx(d,'本文针对玩家只能获知当天天气、无法预知未来天气的条件，构造有限时域自适应鲁棒动态决策模型。第三关完整处理2¹⁰种无沙暴天气情景；第四关以沙暴预算Γ参数化“较少沙暴”，并用题内30日天气估计的一阶Markov链作样本外评价。全部数值均由正式Python程序运行获得。',False);d.add_page_break()
 hd(d,'1 针对第二问的模型构建',1);hd(d,'1.1 建模目标与基本原理',2)
 tx(d,'第二问的实质不是预测一条未来天气，而是在环境信息逐日揭示时构造可执行的反馈策略。若沿用第一问的固定路线和采购量，便隐含使用了尚未公开的未来天气。因此采用自适应鲁棒动态规划：把一天视为一个阶段，把位置、库存和现金作为状态；玩家观察当天天气后行动，对所有符合题意的未来天气分支取最不利值，再选择使该最不利值最大的行动。Bellman最优性原理保证当前最优决策只需依赖当前状态与后续最优价值。')
 tx(d,'建模步骤为：地图与参数结构化→统一日内事件顺序→定义状态和合法动作→建立资源及资金转移→构造非前视策略集→第三关完整情景树求解→第四关Γ预算安全策略求解→Markov辅助评价→全情景、压力情景、Oracle-Regret和Monte Carlo联合检验。概率模型只负责评价，不进入鲁棒主目标，避免用29次天气转移样本制造“精确概率已知”的假象。')
 hd(d,'1.2 状态、决策变量与信息约束',2)
 tx(d,'设地图为无向图G=(V,E)，第t日天气ωₜ∈{S,H,X}分别表示晴朗、高温和沙暴。观察天气后的状态为sₜ=(iₜ,Wₜ,Fₜ,Cₜ)，其中iₜ为位置，Wₜ、Fₜ为水和食物箱数，Cₜ为现金。第四关增加剩余沙暴预算bₜ=Γ−qₜ。决策包括移动xₜ,ᵢⱼ、停留yₜ、挖矿mₜ以及购买量bᵂₜ、bᶠₜ。')
 eq(d,'sₜ=(iₜ,Wₜ,Fₜ,Cₜ),  bₜ=Γ−∑ₖ₌₁ᵗ1(ωₖ=X)',1)
 tx(d,'在线策略πₜ是可见历史hₜ与当天天气到动作的映射。若两种完整天气在第t日前具有相同可见前缀，则必须采取相同行动，即满足非前视约束，从数学上阻止模型偷看未来。')
 eq(d,'hₜ(ω)=hₜ(ω′)  ⇒  πₜ(hₜ,ωₜ)=πₜ(hₜ,ω′ₜ)',2)
 hd(d,'1.3 目标函数与约束体系',2)
 tx(d,'玩家在τ≤T日到达终点后立即结束，剩余资源按基准价格一半回收。鲁棒模型在满足非前视性的策略中最大化最坏天气下的终端财富；鲁棒值并列时，再以经验Markov模型的期望财富作字典序判别，不设置主观风险权重。')
 eq(d,'Z(π,ω)=Cτ+0.5pᵂWτ+0.5pᶠFτ',3);eq(d,'max_{π∈Πᴺᴬ} min_{ω∈Ω} Z(π,ω)',4)
 tx(d,'每日只能执行移动、停留或挖矿中的一个主动作。令λₜ分别取1、2、3表示三类动作消耗倍率，cᵂ(ωₜ)、cᶠ(ωₜ)为天气基础消耗。第四关村庄按基准价两倍补给，第0日在起点按基准价采购。')
 eq(d,'Wₜ₊₁=Wₜ+bᵂₜ−λₜcᵂ(ωₜ),   Fₜ₊₁=Fₜ+bᶠₜ−λₜcᶠ(ωₜ)',5)
 eq(d,'Cₜ₊₁=Cₜ−ρₜ(pᵂbᵂₜ+pᶠbᶠₜ)+R·mₜ,   ρₜ∈{1,2}',6)
 eq(d,'3Wₜ+2Fₜ≤1200,  Wₜ,Fₜ∈Z₊,  Cₜ≥0',7);eq(d,'∑_{j∈N(iₜ)}xₜ,ᵢₜⱼ+yₜ+mₜ=1',8)
 tx(d,'若ωₜ=X，所有移动变量为0；只有日初已在矿山才能挖矿，移动到矿山当日不能同时挖矿。到达村庄前的移动消耗必须由移动前库存支付。到达终点后状态吸收，不允许离开终点继续获利。')
 hd(d,'1.4 第三关：完整情景自适应鲁棒模型',2)
 tx(d,'第三关截止期10日且明确无沙暴，故Ω₃={S,H}¹⁰，共1024条完整情景。令Vₜ(s,ω)表示已知第t日天气时，从状态s出发可保证的最大最坏终端财富。当日先选动作，下一日天气随后揭示，得到鲁棒Bellman方程：')
 eq(d,'Vₜ(s,ω)=max_{a∈A(s,ω)} min_{ω′∈{S,H}} Vₜ₊₁(T(s,a,ω),ω′)',9)
 tx(d,'程序将递推等价展开为非前视天气情景树混合整数线性规划。相同天气前缀共用状态与动作变量，所有终端叶节点受共同财富下界η约束并最大化η。该形式既保留动态反馈，又可由HiGHS提供MIP gap=0的全局最优性证据。')
 hd(d,'1.5 第四关：Γ预算不确定集与安全策略',2)
 tx(d,'第四关只给出“30日内较少出现沙暴”，无法识别唯一概率。定义预算不确定集Ω₄(Γ)，其中Γ表示30日中允许出现的最大沙暴日数。若剩余预算b>0，下一日允许S、H、X；若b=0，则只允许S、H。理论递推为：')
 eq(d,'Ω₄(Γ)={ω∈{S,H,X}³⁰:∑ₜ1(ωₜ=X)≤Γ}',10)
 eq(d,'Vₜ(s,ω,b)=max_{a∈A(s,ω)} min_{ω′∈W(b)}Vₜ₊₁(T(s,a,ω),ω′,b−1(ω′=X))',11)
 eq(d,'W(b)={S,H,X}, b>0;   W(0)={S,H}',12)
 tx(d,'30日、村庄购买和矿山收益会造成大规模状态爆炸。当前正式实现先构造可证明安全下界：取起点1至终点25的8步最短路；非沙暴日沿最短路移动，沙暴日停留；按全部非沙暴移动日为高温且恰有Γ个沙暴预购资源。该策略不声称是含村庄和矿山收益的完整全局最优鲁棒前沿，但能严格证明在Ω₄(Γ)内可行。')
 eq(d,'W₀(Γ)=8×2×max{3,9}+10Γ=144+10Γ',13);eq(d,'F₀(Γ)=8×2×max{4,9}+10Γ=144+10Γ',14)
 eq(d,'Z̲(Γ)=10000−5W₀(Γ)−10F₀(Γ)=7840−150Γ',15)
 tx(d,'由8+Γ≤30可保证时限；由3W₀+2F₀≤1200可得Γ≤9，故Γ=0,…,9为负重可行范围。题内历史序列含6个沙暴日，因此选择Γ=6作为主要推荐压力等级。')
 hd(d,'1.6 Markov辅助模型与评价指标',2)
 tx(d,'题内30日天气含晴朗9日、高温15日、沙暴6日，共29次相邻转移。以频数极大似然估计一阶Markov矩阵P̂。该矩阵仅用于生成统一测试集和计算名义指标，不改变鲁棒策略。')
 eq(d,'P̂=[[2/9,5/9,2/9],[5/14,6/14,3/14],[2/6,3/6,1/6]]',16)
 eq(d,'SuccessRate=N⁻¹∑ₖ1(Successₖ=1)',17);eq(d,'Regretₖ=Z*ₖ−Zπₖ',18)
 tx(d,'其中Z*ₖ为同一天气下完全信息Oracle财富上界。评价采用成功率、成功样本平均财富、最小财富、5%分位财富、平均到达日和Regret，分别对应可行性、经济性、尾部安全性、时效性和相对信息上界的损失。')
 d.add_page_break();hd(d,'2 针对第二问的求解方法',1);hd(d,'2.1 求解工具与算法流程',2)
 tx(d,'全部模型用Python实现。NumPy负责矩阵计算和随机数生成；SciPy milp调用HiGHS求解第三关非前视情景树MILP；collections.deque完成BFS；字典和Pareto标签用于状态压缩；独立transition模块统一执行移动、停留、挖矿、购买与终点回收规则。Monte Carlo固定随机种子20200816，保证结果可复现。')
 tx(d,'第三关步骤为：生成1024个S/H叶情景→按公共前缀构造2046个天气树节点→建立状态、行动和终端下界约束→HiGHS分支定界→检查最优状态与MIP gap→导出在线策略树→逐情景模拟→Oracle DP对照。收敛判据为求解器返回全局最优且MIP gap=0。')
 tx(d,'第四关步骤为：BFS获得8步最短路→扫描Γ并由式(13)—(15)计算采购与保证财富→构造“Γ个沙暴前置、随后全部高温移动”的压力情景→逐日重放→由P̂生成10000组30日天气→同样本比较三种策略→扫描Γ=0,…,9→扰动沙暴概率倍率0.7—1.3。')
 hd(d,'2.2 第三关关键求解结果',2);q3=SUMMARY['第三关']
 tb(d,['指标','结果'],[['初始采购','水54箱、食物54箱，采购后现金9190元'],['鲁棒最坏终端财富','9190元'],['名义Markov期望财富',f"{q3['名义Markov期望财富']:.2f}元"],['全情景平均财富','9310元'],['5%分位财富','9190元'],['到达日','1024个情景均在第3日到达'],['MILP规模','196419变量、94095约束、2046天气树节点'],['求解状态','HiGHS全局最优，MIP gap=0']],9)
 tx(d,'初购54箱水和54箱食物总重量270 kg，成本810元。在线策略沿三步最短路径迅速到达终点，不依赖矿山收益。最坏情形无剩余资源，终端财富等于采购后现金9190元；温和天气产生资源回收，使平均财富升至9310元。名义Markov期望9282.38元低于1024情景等权均值，是因为经验矩阵对高温分支赋予更高权重。')
 hd(d,'2.3 第四关关键求解结果',2);q4=SUMMARY['第四关']['Gamma安全下界']
 tb(d,['Γ','初购水','初购食物','最迟保证到达日','保证财富下界/元'],[[x['Gamma'],x['初购水'],x['初购食物'],x['最迟保证到达日'],f"{x['保证财富下界']:.0f}"] for x in q4],8.5)
 tx(d,'Γ每增加1，水和食物分别增加10箱，初始成本增加150元，保证财富线性下降150元，最迟保证到达日延后1日。Γ=6时初购水、食物各204箱，总重量1020 kg，采购后现金及保证财富下界均为6940元，最迟第14日到达，仍保留16天时间冗余和180 kg负重冗余。')
 hd(d,'3 针对第二问的结果分析',1);hd(d,'3.1 基础分析',2)
 tx(d,'第三关1024种允许天气全部成功且均在第3日到达，说明天气只改变资源余量和回收财富，不改变最短到达结构。最坏财富9190元、均值9310元、5%分位9190元，下尾被鲁棒目标稳定托住。由于挖矿收益仅200元/日，而绕行和挖矿会产生额外消耗，快速到终点优于追逐矿山收益。')
 tx(d,'第四关Γ=6安全策略在10000次Markov样本中成功9884次，成功率98.84%；成功样本平均财富7489.81元，最小财富6940元，5%分位7250元，平均第10.01日到达。样本均值高于6940元保证下界，是因为实际非沙暴日并非全部高温，且多数样本在完成8次移动前的沙暴少于6日。')
 hd(d,'3.2 与第一问及简单策略的统一对比',2)
 tb(d,['策略','成功率','平均财富','最小财富','Q05','平均到达日','平均Regret'],[[r['strategy'],f"{float(r['success_rate']):.2%}",f"{float(r['mean_wealth']):.2f}",f"{float(r['minimum_wealth']):.0f}",f"{float(r['q05_wealth']):.0f}",f"{float(r['mean_arrival_day']):.2f}",f"{float(r['mean_regret']):.2f}"] for r in MET],7.5)
 tx(d,'同一批10000组天气下，Γ=6鲁棒方案成功率98.84%，比第一问已知天气固定方案56.87%高41.97个百分点，比Γ=2低保护方案66.38%高32.46个百分点。第一问方案成功条件下平均财富8127.39元，表面高于鲁棒方案，但其43.13%的样本无法完成任务；该条件均值排除了失败样本，不能单独判断策略优劣。')
 tx(d,'第一问方案针对一条事先已知天气优化，安全库存只覆盖特定排列；当天气沙暴次数或位置改变时，固定采购无法反馈修正。鲁棒方案以额外安全库存将失败概率由43.13%降至1.16%，相对降低约97.31%。其优势并非追求每个成功样本的最高财富，而是在未知未来条件下显著提高任务完成可靠性，直接回应第二问核心要求。')
 hd(d,'3.3 Γ灵敏性与安全—收益前沿',2)
 tb(d,['Γ','采购后现金','成功率','平均财富','Q05','平均Regret'],[[r['Gamma'],r['初购后现金'],f"{float(r['成功率']):.2%}",f"{float(r['成功样本平均财富']):.2f}",f"{float(r['5%分位财富']):.0f}",f"{float(r['平均Regret']):.2f}"] for r in GAM],8)
 tx(d,'Γ从0增至9，成功率由14.27%单调升至99.98%，但成功样本平均财富由8090.04元降至7260.20元，平均Regret由250.04元升至770.20元，形成明确的安全—收益权衡。Γ=0至4每增加一级带来较大成功率增益；Γ=5至6仍由97.21%提高到98.84%；此后边际增益迅速缩小，Γ=6至9仅再提高1.14个百分点，却额外占用450元。结合历史样本含6个沙暴日，Γ=6是兼顾数据锚定、成功率和经济性的折中方案。')
 hd(d,'3.4 天气分布扰动灵敏性',2)
 tb(d,['沙暴倍率','成功率','平均财富','Q05','平均到达日','平均Regret'],[[r['沙暴概率倍率'],f"{float(r['成功率']):.2%}",f"{float(r['成功样本平均财富']):.2f}",f"{float(r['5%分位财富']):.2f}",f"{float(r['平均到达日']):.2f}",f"{float(r['平均Regret']):.2f}"] for r in STO],8)
 tx(d,'将Markov矩阵沙暴列权重乘0.7、0.85、1.0、1.15、1.3并归一化后重抽样。沙暴倍率提高30%时，成功率仍为96.70%，平均财富7452.60元、Q05为7175元，平均到达日延至10.51日。变化方向符合机制：沙暴越多，停留、消耗和到达时间上升。主要结论在较大分布扰动下未反转，说明策略不依赖Markov单点估计。')
 d.add_page_break();hd(d,'4 针对第二问的模型检验',1);hd(d,'4.1 规则一致性和非前视检验',2)
 tx(d,'首先由独立模拟器逐日重算：读取日初状态→检查购买节点和资金→检查天气下行动合法性→按倍率扣除资源→核验库存非负与3W+2F≤1200→更新现金→检查到达日。第三关还比较相同天气历史前缀的动作；任何动作不一致均判为非前视失败。实际资源、资金、沙暴禁行、邻接、终点和非前视检验全部通过。')
 hd(d,'4.2 第三关全情景与Oracle-Regret检验',2)
 tb(d,['检验量','结果','判据'],[['全情景数','1024','2¹⁰无遗漏、无重复'],['成功/失败','1024/0','鲁棒可行须零失败'],['最坏财富','9190元','与优化鲁棒值一致'],['平均/中位/最大Regret','134.45/145/240元','非负且有限'],['最小Regret','0元','在线值不超过Oracle'],['到达分布','第3日：1024次','全部不超过第10日'],['求解器','最优，MIP gap=0','全局最优性判据']],8.5)
 tx(d,'最坏情景逐日重放财富恰为9190元，与MILP目标相等；最小Regret为0且无负Regret，说明Oracle和在线策略口径一致。平均Regret134.45元仅占约万元初始资金的1.34%，最大Regret240元，表明在不能预知未来的现实信息约束下，策略与先知上界差距有限。')
 hd(d,'4.3 第四关压力测试与Monte Carlo检验',2)
 tx(d,'对Γ=0,…,6分别构造Γ个沙暴前置、随后8个高温日完成最短路的压力情景。该排列使沙暴等待和最昂贵的非沙暴移动同时发生。全部情景均在第8+Γ日到达，终端财富严格等于式(15)下界，证明预算语义、资源计算和截止期判断一致。')
 tx(d,'Monte Carlo以固定种子生成10000条30日天气，让三种策略在同一测试集和规则模拟器运行。鲁棒策略98.84%的成功率显著优于两组基线，且沙暴概率提高30%后仍保持96.70%。压力测试覆盖预算集内人为最坏排列，随机模拟考察经验分布下典型性能，两者互补，避免仅凭随机样本漏掉尾部风险。')
 hd(d,'4.4 检验结论与证据边界',2)
 tx(d,'综合全情景枚举、MIP零间隙、独立规则重放、Oracle-Regret、预算压力测试、统一Monte Carlo和参数扰动，可确认第三关策略具有全局最优鲁棒性，第四关Γ=6方案具有预算内可行保证和良好样本外成功率。必须强调：第四关数值是安全策略下界及统计评价，不能据此宣称已求得包含村庄补给、矿山收益的完整全局最优鲁棒策略。')
 hd(d,'5 针对第二问的模型评价',1);hd(d,'5.1 模型优点',2)
 bl(d,'（1）信息结构严格。','非前视约束确保第t日动作只依赖截至当日的信息；第三关1024情景前缀检查全部通过，避免固定路径模型的信息泄漏。')
 bl(d,'（2）主模型不依赖伪概率。','第三关直接覆盖2¹⁰个情景；第四关用Γ表达“较少沙暴”，Markov只作评价，29次历史转移误差不支配主决策。')
 bl(d,'（3）最优性和可靠性证据充分。','第三关196419变量、94095约束的MILP达到MIP gap=0，1024情景零失败，最坏财富9190元，平均Regret仅134.45元。')
 bl(d,'（4）风险改善可量化。','第四关Γ=6成功率98.84%，较第一问固定方案提高41.97个百分点，失败概率相对下降约97.31%；沙暴概率提高30%后仍成功96.70%。')
 bl(d,'（5）可解释且可扩展。','Γ每增加1对应水、食物各增加10箱、保证财富减少150元；统一transition模块可扩展到其他地图、价格和截止期。')
 hd(d,'5.2 模型缺点与改进方向',2)
 bl(d,'（1）第四关当前实现偏保守。','安全策略未利用村庄14和矿山18；虽成功率98.84%，但平均Regret为549.81元。后续可实现带预算状态和Pareto剪枝的完整鲁棒DP。')
 bl(d,'（2）Markov评价样本较小。','转移矩阵仅由29次相邻转移估计，一阶假设不能刻画长记忆。虽已进行0.7—1.3倍沙暴扰动，更充分数据仍可提高外推可信度。')
 bl(d,'（3）状态维数增长较快。','第三关情景树已有196419个变量；第四关若同时精确展开30日、三天气、村庄、矿山和整数库存，规模显著膨胀。可考虑功能节点压缩、列生成或近似DP。')
 hd(d,'5.3 对问题的直接回答',2)
 tx(d,'第三关推荐起点购买水54箱、食物54箱，并依据每天揭示的晴朗或高温执行非前视策略；1024种允许天气均于第3日到达，最坏财富9190元。第四关不应给出依赖人为概率的单一路线，而应以Γ展示安全—收益前沿。结合历史天气含6个沙暴日和Monte Carlo拐点，推荐Γ=6保底方案：初购水、食物各204箱，预算集内最迟第14日到达，保证财富下界6940元；10000次样本外模拟成功率98.84%，显著优于第一问固定方案。')
 hd(d,'参考文献',1)
 for s in ['[1] 全国大学生数学建模竞赛组委会. 2020年高教社杯全国大学生数学建模竞赛B题：穿越沙漠.','[2] Ramani S, Ghate A. Robust Markov Decision Processes with Data-Driven, Distance-Based Ambiguity Sets. SIAM Journal on Optimization, 2022, 32(2): 989–1017.','[3] Goyal V, Grand-Clément J. Robust Markov Decision Processes: Beyond Rectangularity. Mathematics of Operations Research, 2023, 48(1).','[4] 本项目第二问建模手册（V2）及配套Python求解、验证和结果文件.']:tx(d,s,False)
 OUT.parent.mkdir(parents=True,exist_ok=True);d.save(OUT);print(OUT)
if __name__=='__main__':build()
