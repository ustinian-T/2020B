from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = PROJECT_ROOT / "求解代码" / "第一问" / "结果输出"
REPORT_PATH = PROJECT_ROOT / "求解报告" / "第一问求解报告.docx"
SUMMARY = json.loads((OUTPUT_DIR / "求解摘要.json").read_text(encoding="utf-8"))
with (OUTPUT_DIR / "灵敏性分析.csv").open(encoding="utf-8-sig", newline="") as handle:
    SENSITIVITY = list(csv.DictReader(handle))


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 98, 108)
LIGHT_FILL = "F4F6F9"
GRAY_FILL = "F2F4F7"


def set_run_font(run, east_asia="宋体", ascii_font="Calibri", size=11, bold=None, color=None, italic=None):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != 9360:
        raise ValueError(f"table widths must sum to 9360, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_three_line_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name, val, size in (
        ("top", "single", "12"), ("bottom", "single", "12"),
        ("left", "nil", "0"), ("right", "nil", "0"),
        ("insideH", "nil", "0"), ("insideV", "nil", "0"),
    ):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), "7F8C99")
        borders.append(edge)
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "7F8C99")
        tc_borders.append(bottom)
        tc_pr.append(tc_borders)


def style_table(table, widths, header_fill=GRAY_FILL, font_size=9.5):
    set_table_geometry(table, widths)
    set_three_line_borders(table)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "1")
        tr_pr.append(cant_split)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        east_asia="黑体" if row_index == 0 else "宋体",
                        size=font_size,
                        bold=row_index == 0,
                    )
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, east_asia="黑体", bold=True, size=11)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def add_formula(doc, text, number=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text + (f"    ({number})" if number else ""))
    set_run_font(run, east_asia="宋体", ascii_font="Cambria Math", size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(
        run,
        east_asia="黑体",
        ascii_font="Calibri",
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color=BLUE if level < 3 else DARK_BLUE,
    )
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size=9.5, bold=True, color=INK)
    return p


def add_key_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table, widths, font_size=font_size)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for level, size, before, after, color in (
        (1, 16, 18, 10, BLUE), (2, 13, 12, 6, BLUE), (3, 12, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("2020B 穿越沙漠｜第一问求解报告")
    set_run_font(run, east_asia="黑体", size=9, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = p.add_run("第 ")
    set_run_font(lead, size=9, color=MUTED)
    add_page_number(p)
    tail = p.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("数学建模求解报告")
    set_run_font(run, east_asia="黑体", size=12, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("穿越沙漠：已知天气条件下的单玩家最优策略")
    set_run_font(run, east_asia="黑体", size=25, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("第一问建模、求解、结果分析与模型检验")
    set_run_font(run, east_asia="黑体", size=15, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(70)
    run = p.add_run("完整原图｜确定性有限期资源约束优化｜独立逐日复算")
    set_run_font(run, size=10.5, italic=True, color=MUTED)
    for label, value in (
        ("覆盖关卡", "第一关（27 节点）与第二关（64 节点）"),
        ("求解目标", "在第 30 天或之前到达终点并最大化终端总财富"),
        ("结果", "第一关 10470 元；第二关 12730 元"),
        ("生成日期", "2026 年 8 月 15 日"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(label + "：")
        set_run_font(r1, east_asia="黑体", size=10.5, bold=True, color=INK)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5)
    doc.add_page_break()


def add_problem_sections(doc):
    add_heading(doc, "1 问题重述", 1)
    add_body(doc, "在应急物流、极端环境运输、野外作业与移动能源管理中，决策者常常需要在有限载荷和有限资金下，穿越具有时间变化风险的区域。此类任务的现实痛点并不只是“怎样走得最短”，而是要同时回答何时移动、何时等待、携带多少补给、是否绕行获取收益以及何时再次采购。若忽略天气造成的日际消耗差异，最短路线可能因库存不足而不可行；若只追求矿山收益，又可能因额外停留和三倍资源消耗使净财富下降。因此，该问题属于典型的时间依赖、多资源耦合的序贯决策问题。")
    add_body(doc, "本问给定单个玩家、两类离散资源和完全已知的 30 天天气序列。玩家第 0 天以初始资金在起点采购水和食物，之后在地图上逐日选择行走、停留或在矿山挖矿；经过村庄时可以按两倍基准价补给，到达终点后剩余资源按半价回收。核心目标是在不晚于截止日期到达终点的前提下，使现金与终点回收价值之和最大。约束包括地图邻接、沙暴禁行、库存非负、负重不超过 1200 kg、现金非负、起点只采购一次、到矿当日不能挖矿以及到达终点立即结束。")
    add_heading(doc, "1.1 子问题一：题目规则与地图的可计算化", 2)
    add_body(doc, "首先需要把地图图形、功能节点、天气序列和资源参数转换为统一的数据结构，并验证邻接表对称、节点编号合法、起终点连通以及天气长度正确。第一关采用经人工核对的 27 节点邻接表；第二关采用 8×8 错位六邻域网格。该步骤决定了后续所有路线是否真实合法，是求解可靠性的输入基础。")
    add_heading(doc, "1.2 子问题二：已知天气下的全局最优决策", 2)
    add_body(doc, "在完整天气信息下，需要联合优化每日位置、行动、库存、采购量和挖矿安排。问题兼具资源约束最短路、途中补给和收益节点三个特征。Himmich、El Hallaoui 与 Soumis（2024）指出，多阶段动态规划能够通过状态空间划分提升资源约束最短路的精确求解效率；Davatgari 等（2024）进一步表明，在途中允许资源补给的约束路径问题中，Bellman 框架下的标签修正方法仍具有良好的适配性。这些研究说明，本题应优先利用最优子结构与状态支配，而非采用缺乏最优性保证的黑箱启发式算法。")
    add_heading(doc, "1.3 子问题三：最优性、规则一致性与稳健性验证", 2)
    add_body(doc, "求解器输出一个数值并不等于模型正确。还需以独立规则引擎逐日复算资金、资源与负重，检验沙暴、矿山和村庄时点；通过小型实例和终点吸收回归测试检查模型边界；再对负重、初始资金、矿山收益和截止日期作单因素扰动，观察最优财富和策略结构是否符合经济逻辑。")

    add_heading(doc, "2 问题分析", 1)
    add_body(doc, "总体上，本问采用“数据结构化—有限期状态建模—精确全局求解—策略回溯—独立复算—灵敏性分析”的总路线。图结构决定可选移动，天气决定当日动作的可行性与消耗，库存和现金决定未来行为能否继续，矿山与村庄则引入收益和补给。由于未来决策只依赖当前日期、位置、库存与现金，而不依赖更早的完整历史，问题具有 Bellman 最优性结构。")
    add_heading(doc, "2.1 问题1分析。地图与参数预处理", 2)
    add_body(doc, "关键矛盾在于地图来自不规则图形，任何一条误录边都可能制造不存在的捷径或切断合法路线。处理时先把每个区域作为节点，只对存在公共边界的区域连边；随后建立无向邻接集合，检查边的唯一性、对称性、无自环和起终点连通性。天气被编码为长度 30 的有序数组，消耗表按“天气×行动倍率”预先计算。逻辑衔接为：原始规则与地图→结构化配置→输入断言→BFS 距离下界→交给优化模型。")
    add_heading(doc, "2.2 问题2分析。联合路线—资源—收益优化", 2)
    add_body(doc, "关键矛盾是短路线省资源，而绕经矿山可能增加现金；起点采购便宜但受负重限制，村庄补给灵活却价格翻倍；天气完全已知又使“今天移动还是等待”与未来高温、沙暴紧密耦合。为此，以天为阶段，以位置、水、食物为状态索引，以现金为状态价值，逐日枚举停留、行走、挖矿和村庄补给。完整模型再等价展开为逐日整数网络流，由 HiGHS 做全局优化并回溯逐日策略。逻辑衔接为：合法初始采购→逐日动作→扣除消耗→增加挖矿收益→村庄日末补给→终点结算→最大化终端财富。")
    add_heading(doc, "2.3 问题3分析。模型检验与稳健性", 2)
    add_body(doc, "关键矛盾是优化器只保证其所接收数学模型的最优性，不能自动发现模型是否遗漏题目规则。因此将 checker 与求解器分离：checker 只读取逐日动作和采购记录，从第 0 天重新计算，不使用求解器内部库存或累计收益。再利用小型可枚举实例验证最优值，用终点吸收案例防止到达后离开，并在参数扰动下重新求解。文字流程图为：求解结果→逐日规则重放→是否全部通过？若否则返回定位约束；若是则进行精确性对照→参数扰动→输出可靠性结论。")

    doc.add_page_break()
    add_heading(doc, "3 模型假设", 1)
    assumptions = [
        ("附件参数与天气在关卡内准确且固定。", "第一问明确给出全时段天气，价格、负重、消耗和收益均为确定性规则。", "将问题限定为确定性优化，不引入缺乏依据的概率或风险系数。"),
        ("一天最多执行一种耗时行动；采购为瞬时行为。", "题目分别定义行走、停留和挖矿的一日消耗，且村庄允许随时采购。", "可把购买统一归并到日末，形成清晰的逐日状态转移。"),
        ("村庄同一时点的多次购买可合并。", "村庄单价固定，无数量折扣、固定费用或购买次数限制。", "用一次非负整数采购量替代无限次零时间购买，不改变成本和可行集合。"),
        ("玩家不主动丢弃或途中转售资源。", "题目没有规定丢弃和途中回收，只有终点半价回收。", "避免加入未定义动作；库存变化只来自采购和消耗，便于守恒检验。"),
        ("到达终点立即结束，剩余资源按半价结算。", "这是题目明确的终止规则，且不能预先假设终点库存一定为零。", "终点成为吸收状态，目标函数保留库存回收价值并禁止到达后继续挖矿。"),
    ]
    add_key_table(doc, ["假设内容", "合理性依据", "对模型的影响"], assumptions, [2500, 3400, 3460], 9.3)

    add_heading(doc, "4 符号说明", 1)
    symbols = [
        ("G=(V,E)", "地图对应的无向图；V 为区域集合，E 为公共边界邻接边集合", "—"),
        ("s,g", "起点与终点节点", "节点编号"),
        ("V_v,V_m", "村庄节点集合、矿山节点集合", "—"),
        ("T,t", "截止日期与当前日序号", "天"),
        ("ω_t", "第 t 天已知天气", "—"),
        ("M,C_0,R", "负重上限、初始资金、矿山日基础收益", "kg，元，元/天"),
        ("p_w,p_f", "水、食物基准价格", "元/箱"),
        ("m_w,m_f", "每箱水、食物质量", "kg/箱"),
        ("c_w(ω),c_f(ω)", "天气 ω 下停留一天的水、食物基础消耗", "箱/天"),
        ("a_t", "第 t 天行动：停留、行走或挖矿", "—"),
        ("κ(a_t)", "行动消耗倍率，分别为 1、2、3", "—"),
        ("i_t,W_t,F_t,C_t", "第 t 日日末位置、水库存、食物库存与现金", "节点，箱，箱，元"),
        ("b^w_t,b^f_t", "第 t 天日末在村庄购买的水、食物量", "箱"),
        ("V_t(i,w,f)", "第 t 日日末状态可达到的最大现金", "元"),
        ("d(i,g)", "节点 i 到终点的最少移动步数", "可移动日"),
        ("Z", "到达终点后的终端总财富", "元"),
    ]
    add_key_table(doc, ["符号", "含义", "单位"], symbols, [1600, 5660, 2100], 9.2)


def add_model(doc):
    doc.add_page_break()
    add_heading(doc, "5 模型建立", 1)
    add_heading(doc, "5.1 图模型与数据预处理", 2)
    add_body(doc, "将地图表示为无向图 G=(V,E)。若区域 i 与 j 共享公共边界，则 (i,j)∈E；仅共享顶点不连边。每次合法移动恰好占用一天，故边权统一为 1。BFS 只计算节点到终点的最少移动步数 d(i,g)，用于可达性检查，不直接替代天气、资源和收益决策。第一关共有 27 个节点、53 条无向边；第二关共有 64 个节点、161 条无向边。")
    add_heading(doc, "5.2 决策变量、状态变量与价值函数", 2)
    add_body(doc, "每日耗时决策 a_t∈{Stay, Move(j), Mine}。若行动为 Move(j)，必须满足 j∈N(i_{t-1})；采购决策 b^w_t、b^f_t 为非负整数。状态采用日末口径 S_t=(t,i_t,W_t,F_t)，现金不作为 DP 索引，而由价值函数保存。")
    add_formula(doc, "V_t(i,w,f)=第 t 日日末处于节点 i、余水 w、余食 f 时可达到的最大现金", "1")
    add_body(doc, "这一降维利用了最优子结构：若两个历史在同日、同位置、同库存汇合，未来可选动作集合完全相同，现金较多者必不劣。因此无需记录完整历史，只需保存最优前驱。")
    add_heading(doc, "5.3 目标函数", 2)
    add_body(doc, "设玩家在第 τ 天到达终点。现金之外，剩余资源按基准价格的一半退回，因此终端目标不能简单写为“最大剩余现金”，也不能预先把终点库存设为零。")
    add_formula(doc, "max Z = C_τ + 0.5 p_w W_τ + 0.5 p_f F_τ", "2")
    add_formula(doc, "s.t.  0≤τ≤T，i_τ=g", "3")
    add_heading(doc, "5.4 初始采购约束", 2)
    add_body(doc, "第 0 天仅在起点按基准价采购一次。水、食物箱数为非负整数，并同时满足预算和负重。返回起点不再产生采购权限。")
    add_formula(doc, "C_0 = 10000 - p_w W_0 - p_f F_0", "4")
    add_formula(doc, "p_wW_0+p_fF_0≤10000，m_wW_0+m_fF_0≤1200，W_0,F_0∈Z_+", "5")
    add_heading(doc, "5.5 行动合法性与消耗", 2)
    add_body(doc, "停留、行走、挖矿的消耗倍率分别为 1、2、3。沙暴不是改变倍率，而是直接禁止行走；挖矿要求当日开始时已位于矿山，因此到达矿山当天不能同时挖矿。")
    add_formula(doc, "κ(a_t)=1(停留)，2(行走)，3(挖矿)", "6")
    add_formula(doc, "ΔW_t=κ(a_t)c_w(ω_t)，ΔF_t=κ(a_t)c_f(ω_t)", "7")
    add_formula(doc, "Move(j)可行⇔j∈N(i_{t-1})且ω_t≠沙暴；Mine可行⇔i_{t-1}∈V_m", "8")
    add_heading(doc, "5.6 资源与现金递推", 2)
    add_body(doc, "每天先完成耗时行动并扣除消耗，再执行日末村庄采购。这样可确保当天刚到村庄时，移动消耗必须由到达前库存承担，避免把“先买后走”错误地混入同一天。")
    add_formula(doc, "W_t=W_{t-1}-κ(a_t)c_w(ω_t)+b^w_t", "9")
    add_formula(doc, "F_t=F_{t-1}-κ(a_t)c_f(ω_t)+b^f_t", "10")
    add_formula(doc, "C_t=C_{t-1}+R·I(a_t=Mine)-2p_wb^w_t-2p_fb^f_t", "11")
    add_formula(doc, "W_t,F_t,C_t≥0，m_wW_t+m_fF_t≤M", "12")
    add_heading(doc, "5.7 村庄补给与终点吸收", 2)
    add_body(doc, "只有日末位置属于村庄时，b^w_t、b^f_t 才可为正；购买后立即检查现金与负重。到达终点后，位置、现金和库存保持不变，不再产生移动、停留消耗、采购或挖矿收益。该吸收约束是模型检验中的重点，因为若允许离开终点，会产生违规的二次挖矿收益。")
    add_formula(doc, "b^w_t=b^f_t=0，若 i_t∉V_v；到达 g 后 a_{t+1}=Finish", "13")
    add_heading(doc, "5.8 Bellman 转移与状态支配", 2)
    add_body(doc, "设从上一日状态 s 经合法行动 a 和日末采购 b 到达 s′，对应现金增量为 r_t(s,a,b)，则 Bellman 方程为：")
    add_formula(doc, "V_t(s′)=max_{(s,a,b)→s′}{V_{t-1}(s)+r_t(s,a,b)}", "14")
    add_body(doc, "在相同日期和节点，若标签 A 的水、食物和现金均不低于标签 B，且至少一项严格更高，则 A 可以模仿 B 的全部后续行动，B 可安全删除。这一 Pareto 支配不引入经验阈值，不改变全局最优。BFS 距离与未来非沙暴日还构成最早到达下界；剩余天数乘基础收益构成乐观收益上界。")
    add_heading(doc, "5.9 等价整数展开与精确求解", 2)
    add_body(doc, "为避免直接枚举约 30×|V|×401×601 个稠密状态，工程求解将上述逐日 Bellman 转移等价展开为网络流整数模型。用 x_{t,i} 表示第 t 日日末是否位于节点 i，用 y_{t,i,j} 表示第 t 天是否沿边 (i,j) 行走，用 u^S_{t,i}、u^M_{t,i} 表示停留和挖矿。每个日初位置恰好选择一个动作，每个日末位置由唯一动作产生；库存、现金和采购量仍按式（9）—（12）递推。")
    add_formula(doc, "Σ_j y_{t,i,j}+u^S_{t,i}+u^M_{t,i}+u^G_t=x_{t-1,i}", "15")
    add_formula(doc, "x_{t,j}=Σ_i y_{t,i,j}+u^S_{t,j}+u^M_{t,j}+u^G_t I(j=g)", "16")
    add_body(doc, "该展开与有限期状态模型在每日动作、资源和资金层面一一对应。SciPy 的 milp 接口调用 HiGHS 分支定界求解，要求相对 MIP 间隙为 0；随后输出的逐日策略由独立 checker 再次复算。")


def daily_rows(level_name):
    result = SUMMARY["关卡结果"][level_name]
    initial = result["初始采购"]
    cash0 = 10000 - 5 * initial["water"] - 10 * initial["food"]
    first = result["逐日策略"][0]
    rows = [[0, "—", first["from_node"], "起点采购", initial["water"], initial["food"], cash0, initial["water"], initial["food"]]]
    for record in result["逐日策略"]:
        rows.append([
            record["day"], record["weather"], record["to_node"], record["action"],
            record["buy_water"], record["buy_food"], record["cash"], record["water"], record["food"],
        ])
    return rows


def add_results(doc):
    add_heading(doc, "6 模型求解与结果分析", 1)
    add_heading(doc, "6.1 求解工具与步骤", 2)
    add_body(doc, "求解代码使用 Python 3.10、NumPy、SciPy 1.15.3 和 HiGHS。首先加载关卡配置并运行地图、天气和参数断言；其次建立逐日位置流、行动、库存、现金和采购变量；然后加入沙暴禁行、矿山、村庄、终点吸收与资源守恒约束；调用 milp 求得零 MIP 间隙的全局最优解；最后按变量值回溯每日动作，并由 checker 从第 0 天独立重算。")
    add_body(doc, "求解流程可写为：关卡配置→图结构校验→建立第 0 天采购变量→逐日位置流和行动约束→库存/现金递推→终点目标→HiGHS 分支定界→收敛判断（MIP gap=0）→策略提取→独立规则复算→输出 CSV、JSON 和 Result.xlsx。")

    first = SUMMARY["关卡结果"]["第一关"]
    second = SUMMARY["关卡结果"]["第二关"]
    add_heading(doc, "6.2 第一关求解结果", 2)
    add_body(doc, f"第一关最优终端财富为 {first['最优终端财富']:.0f} 元，第 {first['到达日期']} 天到达终点。第 0 天购买 {first['初始采购']['water']} 箱水和 {first['初始采购']['food']} 箱食物，采购后现金为 5780 元，初始负重恰为 1200 kg。最优路线依次经过村庄 15 和矿山 12，在矿山有效挖矿 7 天；第 8 天和第 21 天在村庄补给，最终水、食物均为 0。")
    add_caption(doc, "表 6-1 第一关逐日最优策略（日末口径）")
    add_key_table(doc, ["日", "天气", "区域", "行动", "买水", "买食", "现金", "余水", "余食"], daily_rows("第一关"), [650, 1050, 800, 1000, 950, 950, 1600, 1180, 1180], 8.2)
    add_body(doc, "第一关的经济逻辑是：利用起点低价资源承担前半程与矿山作业，在沙暴日不能移动时根据库存选择停留或挖矿；矿山收益累积后回到村庄补足最后一段所需的最少资源。第二次村庄采购后，资源恰好覆盖到终点的三次移动，因此终点库存为零。")

    add_heading(doc, "6.3 第二关求解结果", 2)
    add_body(doc, f"第二关最优终端财富为 {second['最优终端财富']:.0f} 元，第 {second['到达日期']} 天到达终点。第 0 天购买 {second['初始采购']['water']} 箱水和 {second['初始采购']['food']} 箱食物，采购后现金为 5300 元。策略先经村庄 39 补给，再到矿山 30 挖矿 6 天；第 19 天回到村庄 39 补给后前往矿山 55，连续挖矿 7 天，最后经 56 到达终点。")
    add_caption(doc, "表 6-2 第二关逐日最优策略（日末口径）")
    add_key_table(doc, ["日", "天气", "区域", "行动", "买水", "买食", "现金", "余水", "余食"], daily_rows("第二关"), [650, 1050, 800, 1000, 950, 950, 1600, 1180, 1180], 8.2)
    add_body(doc, "第二关的地图包含两个村庄和两个矿山，较大的空间使“先补给—矿山 30—回村—矿山 55—终点”的链式策略优于只访问一个收益节点。到达日恰为截止日，说明时间约束在第二关为紧约束；最终库存为零，说明最后阶段采购与移动消耗实现精确匹配。")
    add_heading(doc, "6.4 策略建议", 2)
    add_body(doc, "在天气完全已知的单玩家场景中，建议遵循四项原则。第一，最短路只能作为移动下界，不能替代库存和收益联合优化。第二，起点资源单价最低，应优先在负重允许范围内携带可跨越前半程的资源，但不应无目的装满某一种资源。第三，沙暴日若已在矿山，是否挖矿应比较 1000 元收益与三倍消耗的机会成本；若不在矿山则只能停留。第四，村庄采购应面向下一补给点或终点做“够用且不过量”的补给，特别是临近终点时避免以双倍价格购入最终只能半价回收的库存。")


def add_validation(doc):
    add_heading(doc, "7 模型检验与灵敏性分析", 1)
    add_heading(doc, "7.1 输入与规则一致性检验", 2)
    checks = [
        ("天气与参数", "30 天天气长度、取值集合、价格、质量、消耗、收益与截止日断言", "通过"),
        ("第一关地图", "27 节点、53 条无向边；对称、无自环、起终点连通", "通过"),
        ("第二关地图", "64 节点、161 条无向边；六邻域模板与边界节点抽查", "通过"),
        ("规则单元测试", "天气消耗、沙暴、村庄采购、矿山、终点回收与资源守恒", "16 项测试全部通过"),
        ("终点吸收回归", "构造矿山只能经终点到达的地图，禁止到达后离开获利", "通过"),
        ("独立逐日复算", "checker 不读取优化器内部累计值，逐日重算现金、库存与负重", "两关均通过"),
        ("整数优化收敛", "HiGHS 最优状态且相对 MIP gap=0", "两关均为 0"),
    ]
    add_key_table(doc, ["检验项目", "检验方法", "实际结果"], checks, [2200, 5000, 2160], 9.1)
    add_body(doc, "检验过程中曾构造出一个可行但非题意最优的异常策略：模型在第 3 天到达终点后继续离开并挖矿。该策略的前三天能通过截断后的规则复算，但第 30 天现金异常增加。通过“终点只能执行 Finish 动作”的回归测试定位并修复后，第一关由 7615 元恢复为 10470 元，说明吸收状态检查对防止虚假收益至关重要。")
    add_heading(doc, "7.2 最优值与规则重放结果", 2)
    add_body(doc, "第一关优化模型规模为 5012 个变量、1924 条约束，HiGHS 在约 10 秒内达到零间隙；第二关为 13779 个变量、4144 条约束，约 23 秒达到零间隙。两关终点现金分别为 10470 元和 12730 元，终点水、食物均为 0。独立 checker 对每一天的邻接、天气、动作、购买、资源、现金、负重和终止条件均未发现违规。")
    add_heading(doc, "7.3 灵敏性分析方法", 2)
    add_body(doc, "采用单因素离散扰动法，每次只改变一个参数并重新求解完整整数模型，而不是在基准解附近做线性外推。第一关考察负重上限 1080/1200/1320 kg、初始资金 9000/10000/11000 元、矿山收益 900/1000/1100 元及截止日 26/28/30 天；第二关重点考察负重和矿山收益。每个场景都要求达到全局最优并通过 checker。")
    sens_rows = [[r["关卡"], r["参数"], r["参数值"], r["最优终端财富"], r["到达日"], r["挖矿天数"]] for r in SENSITIVITY]
    doc.add_page_break()
    add_caption(doc, "表 7-2 参数扰动下的精确重求解结果")
    add_key_table(doc, ["关卡", "扰动参数", "参数值", "最优财富", "到达日", "挖矿天数"], sens_rows, [1200, 1800, 1200, 1800, 1500, 1860], 8.8)
    add_heading(doc, "7.4 灵敏性结果解释", 2)
    add_body(doc, "负重上限对两关均呈显著正向影响。第一关负重由 1200 kg 下调至 1080 kg，最优财富由 10470 元降至 9670 元；上调至 1320 kg 后增至 11270 元。第二关相应由 12130 元上升到 13365 元。这说明容量不仅减少补给成本，还改变可持续挖矿天数和跨补给点运输能力。")
    add_body(doc, "矿山收益的影响近似分段线性但斜率由最优挖矿天数决定。第一关收益从 1000 元降至 900 元时，最优财富降至 9770 元；升至 1100 元时增至 11180 元。第二关分别为 11430 元和 14030 元，变化幅度更大，原因是基准策略共挖矿 13 天。")
    add_body(doc, "第一关初始资金在 9000—11000 元区间内变化时，最优财富分别为 9470、10470、11470 元，呈一比一传导，说明该区间内路线和采购结构未因资金约束发生切换。截止日缩短到 26 天或 28 天，最优财富仍为 10470 元，因为基准策略第 24 天已到达；因此第一关在该区间对截止日具有稳定性，而第二关第 30 天到达，时间约束明显更紧。")

    add_heading(doc, "8 模型评价与使用边界", 1)
    add_heading(doc, "8.1 优点", 2)
    add_body(doc, "模型将图结构、逐日天气、两类库存、现金、采购和矿山收益统一在同一有限期框架中，所有约束都直接来自题目规则，不需要惩罚系数或启发式参数。整数优化达到零间隙，独立 checker 又提供了与求解器解耦的规则证据；配置化设计使两关共用同一套规则和代码。")
    add_heading(doc, "8.2 局限性", 2)
    add_body(doc, "模型依赖天气完全已知，不能直接回答未知天气下的在线决策；变量规模会随天数、地图和玩家数量增加。灵敏性分析采用有限离散点，只能揭示局部区间内的策略切换；若用于更大实例，可进一步实现稀疏标签 DP、Pareto 前沿数据结构或分解算法。")
    add_heading(doc, "8.3 适用边界", 2)
    add_body(doc, "本报告只讨论单玩家、完全已知天气和题设固定价格规则。若加入多人相互影响、随机天气、价格波动或途中资源共享，需要扩展为随机动态规划、鲁棒优化或博弈模型，不能直接沿用本问的确定性最优策略。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] 2020 年高教社杯全国大学生数学建模竞赛 B 题《穿越沙漠》及附件。",
        "[2] Himmich I, El Hallaoui I, Soumis F. A multiphase dynamic programming algorithm for the shortest path problem with resource constraints. European Journal of Operational Research, 2024, 315(2): 470-483. DOI: 10.1016/j.ejor.2023.11.047.",
        "[3] Davatgari A, Mohammadi M Y, Cokyasar T, et al. A Label-correcting Algorithm for Constrained One-to-Many K-shortest Path Problem with Replenishment. Procedia Computer Science, 2024, 238: 369-376. DOI: 10.1016/j.procs.2024.06.037.",
        "[4] 本项目《第一问建模手册：已知天气单玩家最优策略》及配套求解代码、测试与结果输出。",
    ]
    for ref in refs:
        add_body(doc, ref)


def build_report():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "穿越沙漠第一问求解报告"
    doc.core_properties.subject = "已知天气条件下的单玩家最优策略"
    doc.core_properties.author = "数学建模训练项目"
    doc.core_properties.keywords = "动态规划, 资源约束路径, 整数优化, 穿越沙漠"
    add_cover(doc)
    add_problem_sections(doc)
    add_model(doc)
    add_results(doc)
    add_validation(doc)
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
